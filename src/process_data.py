import os
import logging
import argparse
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
import fitz
from tenacity import retry, stop_after_attempt, wait_exponential

# === TẢI CÁC BIẾN MÔI TRƯỜNG ===
load_dotenv()

# === THIẾT LẬP LOGGING ===
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# === LẤY API KEY TỪ BIẾN MÔI TRƯỜNG ===
XAI_API_KEY_FOR_GENT_CONTEXT = os.getenv("XAI_API_KEY_FOR_GENT_CONTEXT")
if not XAI_API_KEY_FOR_GENT_CONTEXT:
    raise ValueError("Không tìm thấy XAI_API_KEY_FOR_GENT_CONTEXT trong biến môi trường!")

# === KHỞI TẠO MÔ HÌNH LLM ===
llm = ChatOpenAI(
    temperature=0,
    model="meta-llama/llama-4-scout-17b-16e-instruct",  
    api_key=XAI_API_KEY_FOR_GENT_CONTEXT,
    base_url="https://api.groq.com/openai/v1"
)

# === SINH NGỮ CẢNH CHO NHIỀU ĐOẠN VĂN BẢN ===
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def invoke_llm(prompt):
    return llm.invoke(prompt)

def generate_contexts(doc_content, chunks):
    """
    Sinh ngữ cảnh cho nhiều đoạn (chunks) từ nội dung tài liệu.
    """
    contexts = []
    for chunk in chunks:
        prompt = f"""
        <document>
        {doc_content}
        </document>

        <chunk>
        {chunk}
        </chunk>

        Hãy cung cấp ngữ cảnh ngắn gọn, mô tả sự liên quan của đoạn này đến toàn bộ tài liệu.
        """
        try:
            response = invoke_llm(prompt)
            contexts.append(response.content.strip())
        except Exception as e:
            logging.error(f"Lỗi khi sinh ngữ cảnh cho đoạn: {e}")
            contexts.append("Không thể sinh ngữ cảnh.")
    return contexts

# === XỬ LÝ FILE DOCX ===
def process_docx(file_path):
    """
    Xử lý một file DOCX, chia thành các đoạn và sinh ngữ cảnh.
    """
    try:
        docx = DocxDocument(file_path)
        doc_content = "\n".join([para.text for para in docx.paragraphs if para.text.strip()])
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=500)
        chunks = text_splitter.split_text(doc_content)
        contexts = generate_contexts(doc_content, chunks)
        all_documents = [
            Document(page_content=f"Đoạn văn bản:\n{chunk}\n\nNgữ cảnh:\n{context}\n")
            for chunk, context in zip(chunks, contexts)
        ]
        logging.info(f"Đã xử lý file DOCX: {file_path} với tổng cộng {len(chunks)} đoạn.")
        return all_documents
    except Exception as e:
        logging.error(f"Lỗi khi xử lý file DOCX {file_path}: {e}")
        return []

# === XỬ LÝ FILE TXT ===
def process_txt(file_path):
    """
    Xử lý một file TXT, chia thành các đoạn và sinh ngữ cảnh.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            doc_content = file.read()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=500)
        chunks = text_splitter.split_text(doc_content)
        contexts = generate_contexts(doc_content, chunks)
        all_documents = [
            Document(page_content=f"Đoạn văn bản:\n{chunk}\n\nNgữ cảnh:\n{context}\n")
            for chunk, context in zip(chunks, contexts)
        ]
        logging.info(f"Đã xử lý file TXT: {file_path} với tổng cộng {len(chunks)} đoạn.")
        return all_documents
    except Exception as e:
        logging.error(f"Lỗi khi xử lý file TXT {file_path}: {e}")
        return []

# === XỬ LÝ FILE PDF ===
def process_pdf(file_path):
    """
    Xử lý một file PDF, chia thành các đoạn và sinh ngữ cảnh.
    """
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        doc_content = " ".join([doc.page_content for doc in documents])
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=1000)
        chunks = text_splitter.split_text(doc_content)
        contexts = generate_contexts(doc_content, chunks)
        all_documents = [
            Document(page_content=f"Đoạn văn bản:\n{chunk}\n\nNgữ cảnh:\n{context}\n")
            for chunk, context in zip(chunks, contexts)
        ]
        logging.info(f"Đã xử lý file PDF: {file_path} với tổng cộng {len(chunks)} đoạn.")
        return all_documents
    except Exception as e:
        logging.error(f"Lỗi khi xử lý file PDF {file_path}: {e}")
        return []

# === KIỂM TRA TÍNH HỢP LỆ CỦA FILE ===
def is_valid_file(file_path):
    """
    Kiểm tra tính hợp lệ của file dựa trên loại file.
    """
    try:
        if file_path.lower().endswith(".docx"):
            docx = DocxDocument(file_path)
            return len([para.text for para in docx.paragraphs if para.text.strip()]) > 0
        elif file_path.lower().endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as file:
                return len(file.read().strip()) > 0
        elif file_path.lower().endswith(".pdf"):
            pdf_document = fitz.open(file_path)
            return pdf_document.page_count > 0
        return False
    except Exception as e:
        logging.warning(f"File không hợp lệ {file_path}: {e}")
        return False

# === DUYỆT VÀ XỬ LÝ TẤT CẢ FILE TRONG THƯ MỤC ===
def process_all_files_in_directory(directory_path):
    """
    Xử lý tất cả file DOCX, TXT, PDF trong thư mục, sinh ngữ cảnh cho từng đoạn.
    """
    all_docs = []
    for root, _, files in os.walk(directory_path):
        for file in files:
            if file.lower().endswith((".docx", ".txt", ".pdf")):
                file_path = os.path.join(root, file)
                if not is_valid_file(file_path):
                    logging.warning(f"Bỏ qua file không hợp lệ: {file_path}")
                    continue
                logging.info(f"Đang xử lý: {file_path}")
                if file.lower().endswith(".docx"):
                    docs = process_docx(file_path)
                elif file.lower().endswith(".txt"):
                    docs = process_txt(file_path)
                elif file.lower().endswith(".pdf"):
                    docs = process_pdf(file_path)
                all_docs.extend(docs)
    logging.info(f"Đã xử lý tổng cộng {len(all_docs)} đoạn từ các file.")
    return all_docs

# === THỰC THI FILE ===
def main():
    parser = argparse.ArgumentParser(description="Process DOCX, TXT, and PDF files and generate contexts.")
    parser.add_argument("--directory", default="D:/DoAnTotNghiep/AIChatFinal/data", help="Directory containing files")
    args = parser.parse_args()
    directory_path = args.directory

    # Xử lý tất cả các file trong thư mục
    processed_docs = process_all_files_in_directory(directory_path)
    
    # Hiển thị thử hai đoạn đầu tiên
    logging.info(f"Đã xử lý tổng cộng {len(processed_docs)} đoạn văn bản từ các file.")
    for doc in processed_docs[:2]:
        print(f"Đoạn văn bản có ngữ cảnh là: \n{doc.page_content}")
        print("=" * 100)

if __name__ == "__main__":
    main()